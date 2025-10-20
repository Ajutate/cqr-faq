import os
from typing import List, Dict
from pathlib import Path
from dotenv import load_dotenv

from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_ollama import OllamaEmbeddings
from langchain_chroma import Chroma
from langchain_ollama import OllamaLLM
from langchain.chains import RetrievalQA
from langchain.prompts import PromptTemplate
from langchain.schema import Document

from pdf_extractor import extract_pdf

load_dotenv()


class RAGService:
    def __init__(self):
        """Initialize RAG service with embeddings and vector store"""
        
        # Initialize Ollama embeddings model (faster and consistent with LLM)
        self.embeddings = OllamaEmbeddings(
            base_url=os.getenv("OLLAMA_BASE_URL", "http://localhost:11434"),
            model=os.getenv("OLLAMA_EMBEDDING_MODEL", "nomic-embed-text")
        )
        
        # Initialize vector store
        self.persist_directory = os.getenv("CHROMA_PERSIST_DIRECTORY", "./chroma_db")
        self.vectorstore = Chroma(
            persist_directory=self.persist_directory,
            embedding_function=self.embeddings,
            collection_name="pdf_documents"
        )
        
        # Initialize Ollama LLM with lower temperature to reduce hallucination
        self.llm = OllamaLLM(
            base_url=os.getenv("OLLAMA_BASE_URL", "http://localhost:11434"),
            model=os.getenv("OLLAMA_MODEL", "qwen3:latest"),
            temperature=0.1,  # Lower temperature = more factual, less creative
        )
        
        # Text splitter for chunking documents - optimized for context preservation
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,  # Larger chunks to preserve more context per embedding
            chunk_overlap=200,  # More overlap to maintain continuity across chunks
            length_function=len,
            separators=["\n\n", "\n", " ", ""]  # Split on natural boundaries
        )
        
        # Custom prompt template - strict instructions to prevent hallucination
        self.prompt_template = """You are a precise document reader. Your job is to extract information EXACTLY as written in the document.

IMPORTANT RULES:
1. Quote numbers and values EXACTLY as they appear in the context
2. Do NOT calculate, estimate, or guess any values
3. Do NOT use information from your training data
4. If the exact information is not in the context below, respond with "I cannot find this specific information in the provided document."
5. When extracting financial data like salary, copy the EXACT figure from the document
6. Pay attention to the source document name in the context - only use information from the relevant document

Context from the document:
{context}

Question: {question}

Answer (quote exact values from context):"""
        
        self.PROMPT = PromptTemplate(
            template=self.prompt_template, 
            input_variables=["context", "question"]
        )
    
    def process_pdf(self, file_path: str) -> Dict:
        """
        Process PDF file and store in vector database
        Uses advanced extraction with table support
        """
        try:
            # Extract PDF using pdfplumber (MIT License - free and open source!)
            extracted_docs = extract_pdf(file_path, method="pdfplumber")
            
            if not extracted_docs:
                raise Exception("No content could be extracted from PDF")
            
            # Convert to LangChain Document format
            documents = []
            for doc in extracted_docs:
                documents.append(Document(
                    page_content=doc["content"],
                    metadata={
                        "source": doc["source"],
                        "page": doc["page"],
                        "method": doc.get("method", "unknown"),
                        "type": doc.get("type", "text")
                    }
                ))
            
            print(f"Extracted {len(documents)} sections from PDF")
            
            # Split documents into chunks
            texts = self.text_splitter.split_documents(documents)
            
            # Add chunk metadata
            for i, text in enumerate(texts):
                text.metadata["chunk"] = i
            
            print(f"Split into {len(texts)} chunks")
            
            # Add to vector store in batches to avoid memory issues with large PDFs
            batch_size = 100
            total_batches = (len(texts) + batch_size - 1) // batch_size
            
            print(f"Adding {len(texts)} chunks to vector store in {total_batches} batches...")
            
            for i in range(0, len(texts), batch_size):
                batch = texts[i:i + batch_size]
                self.vectorstore.add_documents(batch)
                batch_num = (i // batch_size) + 1
                print(f"✓ Batch {batch_num}/{total_batches} added ({len(batch)} chunks)")
            
            print(f"✅ All chunks successfully added to database")
            
            return {
                "status": "success",
                "chunks": len(texts),
                "sections": len(documents),
                "filename": Path(file_path).name
            }
        
        except Exception as e:
            raise Exception(f"Error processing PDF: {str(e)}")
    
    def process_document(self, file_path: str, category: str = None, sub_category: str = None) -> Dict:
        """
        Process document file (PDF, DOC, DOCX, TXT) and store in vector database with optional categorization
        
        Args:
            file_path: Path to the document file
            category: Optional category (e.g., "Home Loan", "Personal Loan")
            sub_category: Optional sub-category (e.g., "Eligibility", "Documentation")
        """
        try:
            file_extension = Path(file_path).suffix.lower()
            
            # Handle different file types
            if file_extension == '.pdf':
                # Extract PDF using pdfplumber
                extracted_docs = extract_pdf(file_path, method="pdfplumber")
                
                if not extracted_docs:
                    raise Exception("No content could be extracted from PDF")
                
                # Convert to LangChain Document format
                documents = []
                for doc in extracted_docs:
                    metadata = {
                        "source": doc["source"],
                        "page": doc["page"],
                        "method": doc.get("method", "unknown"),
                        "type": doc.get("type", "text")
                    }
                    
                    # Add category metadata if provided
                    if category:
                        metadata["category"] = category
                    if sub_category:
                        metadata["sub_category"] = sub_category
                    
                    documents.append(Document(
                        page_content=doc["content"],
                        metadata=metadata
                    ))
            
            elif file_extension in ['.doc', '.docx']:
                # Extract DOC/DOCX using python-docx
                try:
                    from docx import Document as DocxDocument
                    doc = DocxDocument(file_path)
                    
                    # Extract all paragraphs
                    content = '\n'.join([para.text for para in doc.paragraphs if para.text.strip()])
                    
                    if not content.strip():
                        raise Exception("No content could be extracted from DOC/DOCX")
                    
                    metadata = {
                        "source": Path(file_path).name,
                        "page": 1,
                        "method": "python-docx",
                        "type": "text"
                    }
                    
                    # Add category metadata if provided
                    if category:
                        metadata["category"] = category
                    if sub_category:
                        metadata["sub_category"] = sub_category
                    
                    documents = [Document(page_content=content, metadata=metadata)]
                
                except ImportError:
                    raise Exception("python-docx is required for DOC/DOCX files. Install it with: pip install python-docx")
            
            elif file_extension == '.txt':
                # Extract TXT file
                with open(file_path, 'r', encoding='utf-8') as f:
                    content = f.read()
                
                if not content.strip():
                    raise Exception("No content could be extracted from TXT file")
                
                metadata = {
                    "source": Path(file_path).name,
                    "page": 1,
                    "method": "text",
                    "type": "text"
                }
                
                # Add category metadata if provided
                if category:
                    metadata["category"] = category
                if sub_category:
                    metadata["sub_category"] = sub_category
                
                documents = [Document(page_content=content, metadata=metadata)]
            
            else:
                raise Exception(f"Unsupported file type: {file_extension}")
            
            print(f"Extracted {len(documents)} sections from {file_extension.upper()} file")
            if category:
                print(f"   Category: {category}")
            if sub_category:
                print(f"   Sub-Category: {sub_category}")
            
            # Split documents into chunks
            texts = self.text_splitter.split_documents(documents)
            
            # Add chunk metadata (category and sub_category are preserved from parent document)
            for i, text in enumerate(texts):
                text.metadata["chunk"] = i
            
            print(f"Split into {len(texts)} chunks")
            
            # Add to vector store in batches
            batch_size = 100
            total_batches = (len(texts) + batch_size - 1) // batch_size
            
            print(f"Adding {len(texts)} chunks to vector store in {total_batches} batches...")
            
            for i in range(0, len(texts), batch_size):
                batch = texts[i:i + batch_size]
                self.vectorstore.add_documents(batch)
                batch_num = (i // batch_size) + 1
                print(f"✓ Batch {batch_num}/{total_batches} added ({len(batch)} chunks)")
            
            print(f"✅ All chunks successfully added to database")
            
            return {
                "status": "success",
                "chunks": len(texts),
                "sections": len(documents),
                "filename": Path(file_path).name
            }
        
        except Exception as e:
            raise Exception(f"Error processing document: {str(e)}")
    
    def query(self, question: str, k: int = 6, document_filter: str = None, score_threshold: float = 0.3) -> Dict:
        """
        Query the RAG system with similarity search for exact results
        
        Args:
            question: The question to ask
            k: Number of chunks to retrieve
            document_filter: Optional filename to filter results (e.g., "payslip.pdf")
            score_threshold: Minimum similarity score (0.0 to 1.0). Higher = more strict. Default 0.3
        """
        try:
            # Build search kwargs for similarity search with threshold
            search_kwargs = {
                "k": k,  # Number of most similar chunks to retrieve
                "score_threshold": score_threshold,  # Minimum similarity score
            }
            
            # Add document filter if specified
            if document_filter:
                search_kwargs["filter"] = {"source": document_filter}
            
            print(f"🔍 Searching for top {k} most similar chunks")
            
            # Use regular similarity search without threshold
            # ChromaDB uses L2 distance (lower = more similar)
            # We'll retrieve top-k and let the LLM filter the relevant ones
            retriever_kwargs = {"k": k}
            if document_filter:
                retriever_kwargs["filter"] = {"source": document_filter}
            
            # Create retrieval QA chain with similarity search
            qa_chain = RetrievalQA.from_chain_type(
                llm=self.llm,
                chain_type="stuff",
                retriever=self.vectorstore.as_retriever(
                    search_type="similarity",  # Use regular similarity search (top-k)
                    search_kwargs=retriever_kwargs
                ),
                chain_type_kwargs={"prompt": self.PROMPT},
                return_source_documents=True
            )
            
            # Get answer using invoke instead of __call__
            result = qa_chain.invoke({"query": question})
            
            # Extract sources with full content for verification
            sources = []
            source_files = set()
            for doc in result.get("source_documents", []):
                source_file = doc.metadata.get("source", "Unknown")
                source_files.add(source_file)
                sources.append({
                    "content": doc.page_content,  # Return FULL content, not truncated
                    "source": source_file,
                    "page": doc.metadata.get("page", "Unknown")
                })
            
            # Add warning if multiple documents were used
            answer = result["result"]
            if len(source_files) > 1 and not document_filter:
                file_list = ", ".join(source_files)
                answer = f"⚠️ Note: Answer drawn from multiple documents ({file_list}). Consider clearing unwanted documents or specifying which document to query.\n\n{answer}"
            
            return {
                "answer": answer,
                "sources": sources,
                "source_files": list(source_files)
            }
        
        except Exception as e:
            raise Exception(f"Error querying RAG: {str(e)}")
    
    def list_documents(self) -> List[str]:
        """
        List all documents in the vector store
        """
        try:
            # Get all documents
            collection = self.vectorstore.get()
            
            # Extract unique sources
            sources = set()
            if collection and 'metadatas' in collection:
                for metadata in collection['metadatas']:
                    if metadata and 'source' in metadata:
                        sources.add(metadata['source'])
            
            return list(sources)
        
        except Exception as e:
            raise Exception(f"Error listing documents: {str(e)}")
    
    def clear_database(self):
        """
        Clear the vector database
        """
        try:
            # Delete and recreate the vector store
            self.vectorstore.delete_collection()
            self.vectorstore = Chroma(
                persist_directory=self.persist_directory,
                embedding_function=self.embeddings,
                collection_name="pdf_documents"
            )
        
        except Exception as e:
            raise Exception(f"Error clearing database: {str(e)}")
    
    def delete_document(self, filename: str) -> Dict:
        """
        Delete a specific document from the vector store
        
        Args:
            filename: Name of the file to delete (e.g., "payslip.pdf")
        """
        try:
            # Get all documents
            collection = self.vectorstore.get()
            
            # Find IDs of documents with matching source
            ids_to_delete = []
            if collection and 'metadatas' in collection and 'ids' in collection:
                for idx, metadata in enumerate(collection['metadatas']):
                    if metadata and metadata.get('source') == filename:
                        ids_to_delete.append(collection['ids'][idx])
            
            if ids_to_delete:
                self.vectorstore.delete(ids=ids_to_delete)
                return {
                    "status": "success",
                    "deleted": len(ids_to_delete),
                    "filename": filename
                }
            else:
                return {
                    "status": "not_found",
                    "message": f"No documents found with filename: {filename}"
                }
        
        except Exception as e:
            raise Exception(f"Error deleting document: {str(e)}")
